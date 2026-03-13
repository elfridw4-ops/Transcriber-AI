from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import shutil
import os
import uuid
import subprocess
import whisper
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import json

app = FastAPI()

# Configuration CORS pour le frontend React
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Structure des dossiers
BASE_DIRS = ["uploads", "audio", "transcriptions", "subtitles", "documents"]
for d in BASE_DIRS:
    os.makedirs(d, exist_ok=True)

# Chargement du modèle Whisper (Base est un bon compromis vitesse/précision)
print("Chargement du modèle Whisper...")
model = whisper.load_model("base")

@app.post("/api/process")
async def process_video(video: UploadFile = File(...)):
    file_id = str(uuid.uuid4())
    video_ext = video.filename.split(".")[-1]
    video_path = f"uploads/{file_id}.{video_ext}"
    audio_path = f"audio/{file_id}.mp3"
    
    # 1. Sauvegarde de la vidéo
    with open(video_path, "wb") as buffer:
        shutil.copyfileobj(video.file, buffer)
    
    try:
        # 2. Extraction Audio avec FFmpeg
        subprocess.run([
            "ffmpeg", "-i", video_path, 
            "-vn", "-acodec", "libmp3lame", 
            "-y", audio_path
        ], check=True, capture_output=True)
        
        # 3. Transcription avec Whisper (Local)
        print(f"Transcription en cours pour {file_id}...")
        result = model.transcribe(audio_path)
        
        # 4. Sauvegarde des résultats
        text_content = result["text"]
        segments = result["segments"]
        
        # Transcription texte brut
        with open(f"transcriptions/{file_id}.txt", "w", encoding="utf-8") as f:
            f.write(text_content)
            
        # Génération SRT
        srt_path = f"subtitles/{file_id}.srt"
        with open(srt_path, "w", encoding="utf-8") as f:
            for i, segment in enumerate(segments):
                start = format_timestamp(segment['start'])
                end = format_timestamp(segment['end'])
                f.write(f"{i+1}\n{start} --> {end}\n{segment['text'].strip()}\n\n")

        # 5. Génération des documents
        # Word
        doc = Document()
        doc.add_heading('Transcription Vidéo', 0)
        doc.add_paragraph(text_content)
        doc.save(f"documents/{file_id}.docx")
        
        # PDF
        pdf_path = f"documents/{file_id}.pdf"
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.drawString(100, 750, "Transcription Vidéo")
        text_object = c.beginText(100, 730)
        text_object.setFont("Helvetica", 10)
        # Gestion simple du retour à la ligne pour le PDF
        lines = [text_content[i:i+90] for i in range(0, len(text_content), 90)]
        for line in lines:
            text_object.textLine(line)
        c.drawText(text_object)
        c.save()

        return {
            "fileId": file_id,
            "text": text_content,
            "language": result.get("language", "unknown")
        }
        
    except Exception as e:
        print(f"Erreur: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/download/{file_id}/{ext}")
async def download_file(file_id: str, ext: str):
    folder = "documents"
    if ext == "srt": folder = "subtitles"
    if ext == "txt": folder = "transcriptions"
    
    file_path = f"{folder}/{file_id}.{ext}"
    if os.path.exists(file_path):
        return FileResponse(file_path, filename=f"transcription_{file_id}.{ext}")
    raise HTTPException(status_code=404, detail="Fichier non trouvé")

def format_timestamp(seconds: float):
    td = float(seconds)
    hours = int(td // 3600)
    minutes = int((td % 3600) // 60)
    secs = int(td % 60)
    milis = int((td % 1) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{milis:03d}"

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
